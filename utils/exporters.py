import os
import csv
import json
import pandas as pd
from docx import Document
# requirements: pip install python-docx pandas openpyxl reportlab

class Exporter:
    @staticmethod
    def save_to_file(data, format_type, filepath):
        """
        data: List of dictionaries or a single string?
        For batch: list of {'filename':..., 'text':..., 'data':...}
        For single: {'text':...}
        """
        if format_type == "txt":
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(data.get('text', ''))
        elif format_type == "jsonl":
            with open(filepath, 'a', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False)
                f.write('\n')
        elif format_type == "md":
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"# OCR Result\n\n{data.get('text', '')}")
        
    @staticmethod
    def export_batch(results, output_dir, formats):
        """
        results: List of dicts {'filename': str, 'text': str, 'json_data': dict}
        formats: List of extensions ['xlsx', 'docx', 'md']
        """
        os.makedirs(output_dir, exist_ok=True)
        
        for fmt in formats:
            if fmt == "xlsx":
                df = pd.DataFrame(results)
                df.to_excel(os.path.join(output_dir, "batch_result.xlsx"))
                
            elif fmt == "csv":
                df = pd.DataFrame(results)
                df.to_csv(os.path.join(output_dir, "batch_result.csv"))
                
            elif fmt == "docx":
                doc = Document()
                doc.add_heading('OCR Batch Results', 0)
                for res in results:
                    doc.add_heading(os.path.basename(res['filename']), level=1)
                    doc.add_paragraph(res['text'])
                    doc.add_page_break()
                doc.save(os.path.join(output_dir, "batch_result.docx"))
                
            elif fmt == "md":
                with open(os.path.join(output_dir, "batch_result.md"), 'w', encoding='utf-8') as f:
                    for res in results:
                        f.write(f"## {os.path.basename(res['filename'])}\n\n")
                        f.write(res['text'])
                        f.write("\n\n---\n\n")

    @staticmethod
    def create_layered_pdf(image_path, text_data, output_path):
        """
        Requires reportlab. 
        Draws invisible text over the image.
        """
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from PIL import Image
        
        img = Image.open(image_path)
        w, h = img.size
        
        c = canvas.Canvas(output_path, pagesize=(w, h))
        c.drawImage(image_path, 0, 0, w, h)
        
        # Mock implementation of drawing text at coordinates
        # text_data would contain [{'details': 'text', 'box': [x,y,w,h]}]
        # This requires parsing the PaddleOCR standard output structure carefully.
        
        c.save()
