import os
import sys
import json
import csv
import re
import threading
from html.parser import HTMLParser
from PySide6.QtCore import QObject, Signal, QRunnable, QThreadPool
from services.ocr_engine import ocr_client
from app_config import config


class HTMLTableParser(HTMLParser):
    """Parse HTML table into 2D list for Excel export"""
    def __init__(self):
        super().__init__()
        self.tables = []
        self.current_table = []
        self.current_row = []
        self.current_cell = ""
        self.in_table = False
        self.in_row = False
        self.in_cell = False
        self.cell_colspan = 1
        self.cell_rowspan = 1
    
    def handle_starttag(self, tag, attrs):
        if tag == "table":
            self.in_table = True
            self.current_table = []
        elif tag == "tr" and self.in_table:
            self.in_row = True
            self.current_row = []
        elif tag in ("td", "th") and self.in_row:
            self.in_cell = True
            self.current_cell = ""
            # Parse colspan/rowspan
            self.cell_colspan = 1
            self.cell_rowspan = 1
            for attr, val in attrs:
                if attr == "colspan":
                    try: self.cell_colspan = int(val)
                    except: pass
                elif attr == "rowspan":
                    try: self.cell_rowspan = int(val)
                    except: pass
    
    def handle_endtag(self, tag):
        if tag in ("td", "th") and self.in_cell:
            self.in_cell = False
            # Add cell with colspan handling
            for _ in range(self.cell_colspan):
                self.current_row.append(self.current_cell.strip())
        elif tag == "tr" and self.in_row:
            self.in_row = False
            if self.current_row:
                self.current_table.append(self.current_row)
        elif tag == "table" and self.in_table:
            self.in_table = False
            if self.current_table:
                self.tables.append(self.current_table)
    
    def handle_data(self, data):
        if self.in_cell:
            self.current_cell += data


def parse_html_tables(html_content):
    """Parse HTML and extract tables as 2D lists"""
    parser = HTMLTableParser()
    try:
        parser.feed(html_content)
        return parser.tables
    except:
        return []


def parse_markdown_table(md_text):
    """Parse markdown tables to list of 2D lists"""
    tables = []
    current_table = []
    in_table = False
    
    for line in md_text.split("\n"):
        line = line.strip()
        
        # Detect table row
        if line.startswith("|") and line.endswith("|"):
            # Skip separator lines like |---|---|
            stripped = line.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
            if not stripped:
                continue
            
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                current_table.append(cells)
                in_table = True
        else:
            # End of table
            if in_table and current_table:
                tables.append(current_table)
                current_table = []
                in_table = False
    
    # Don't forget last table
    if current_table:
        tables.append(current_table)
    
    return tables if tables else None


def parse_all_markdown_tables(md_text):
    """Extract all tables from markdown content for Excel export."""
    tables = parse_markdown_table(md_text)
    if not tables:
        return None
    return tables


class WorkerSignals(QObject):
    progress_update = Signal(int, int, int)
    status_update = Signal(int, str)
    result_update = Signal(int, str)
    finished = Signal()
    stopped = Signal()  # Emitted when user stops the task
    error_occurred = Signal(int, str)


class BatchWorker(QRunnable):
    def __init__(self, mode, files, model_override, output_dir, export_fmt, max_workers=5, row_indices=None):
        super().__init__()
        self.mode = mode
        self.files = files
        self.model_override = model_override
        self.output_dir = output_dir
        self.export_fmt = export_fmt
        self.max_workers = max_workers
        # row_indices maps internal file indices to actual table row indices
        # If not provided, assume 1:1 mapping (0, 1, 2, ...)
        self.row_indices = row_indices if row_indices else list(range(len(files)))
        self.signals = WorkerSignals()
        self._stop_event = threading.Event()  # Use Event for faster response
        self._completed_count = 0
        self._lock = threading.Lock()

    def stop(self):
        """Request stop - sets event flag immediately"""
        self._stop_event.set()

    def process_single_file(self, index, filepath):
        """Process a single file - runs in thread pool
        
        index: internal index in the files list
        Returns: (index, row_index, status, result)
        
        CRITICAL: Once stop_event is set by user, this method MUST NOT emit any
        signals to UI. This prevents background threads from overwriting the
        "用户已终止任务" status that was set by the main thread.
        """
        # Get the actual table row index for UI updates
        row_index = self.row_indices[index] if index < len(self.row_indices) else index
        
        # Check if stopped before starting - don't waste time on new tasks
        if self._stop_event.is_set():
            # Don't emit anything - main thread already set the status
            return index, row_index, None, "用户已终止任务"

        try:
            # Only update UI if not stopped
            if not self._stop_event.is_set():
                self.signals.status_update.emit(row_index, "处理中...")
                self.signals.progress_update.emit(row_index, 50, 100)

            # Check before reading file
            if self._stop_event.is_set():
                return index, row_index, None, "用户已终止任务"

            with open(filepath, "rb") as f:
                file_bytes = f.read()

            ext = os.path.splitext(filepath)[1].lower()
            is_document = ext in ['.pdf', '.xps', '.epub', '.mobi', '.fb2', '.cbz']
            file_type = 0 if is_document else 1

            # Check before API call
            if self._stop_event.is_set():
                return index, row_index, None, "用户已终止任务"

            # Make API call
            res = ocr_client.ocr_file(file_bytes, file_type=file_type, 
                                       model_override=self.model_override,
                                       stop_event=self._stop_event)

            # ============== CRITICAL SECTION ==============
            # After API call returns, check stop_event FIRST
            # If stopped, DO NOT update UI - just return silently
            # This prevents overwriting "用户已终止任务" set by main thread
            if self._stop_event.is_set():
                return index, row_index, None, "用户已终止任务"
            # ================================================

            # Handle API errors (user termination handled above)
            if "error" in res:
                err_msg = res['error']
                # Double-check stop_event before updating UI
                if self._stop_event.is_set():
                    return index, row_index, None, "用户已终止任务"
                if 'raw_response' in res:
                    err_msg += f"\n{str(res['raw_response'])[:200]}..."
                self.signals.status_update.emit(row_index, "失败")
                self.signals.progress_update.emit(row_index, 100, 100)
                self.signals.result_update.emit(row_index, err_msg)
                return index, row_index, "失败", err_msg

            # Get text content
            text_lines = []
            if "words_result" in res and res["words_result"]:
                for item in res["words_result"]:
                    word = item.get("words", "")
                    if isinstance(word, str):
                        text_lines.append(word)

            markdown_text = res.get("markdown", "")
            tables_html = res.get("tables_html", [])
            layout_data = res.get("layout_data", [])
            raw_result = res.get("raw_result", {})

            full_text = "\n".join(text_lines) if text_lines else markdown_text

            if full_text or tables_html:
                # Save results
                export_result = self.save_results(filepath, full_text, markdown_text, tables_html, layout_data, raw_result)
                
                # Check stop_event before updating UI
                if self._stop_event.is_set():
                    # File was saved but user stopped - don't update UI
                    # The work is done but we respect user's stop request
                    return index, row_index, None, "用户已终止任务"
                
                # Update UI with success
                self.signals.status_update.emit(row_index, "成功")
                self.signals.progress_update.emit(row_index, 100, 100)
                self.signals.result_update.emit(row_index, export_result)
                return index, row_index, "成功", export_result
            else:
                # Check stop_event before updating UI
                if self._stop_event.is_set():
                    return index, row_index, None, "用户已终止任务"
                self.signals.status_update.emit(row_index, "无内容")
                self.signals.progress_update.emit(row_index, 100, 100)
                self.signals.result_update.emit(row_index, "未识别到文字")
                return index, row_index, "无内容", "未识别到文字"

        except Exception as e:
            # Always check stop_event before any UI update
            if self._stop_event.is_set():
                return index, row_index, None, "用户已终止任务"
            self.signals.status_update.emit(row_index, "错误")
            self.signals.progress_update.emit(row_index, 100, 100)
            self.signals.result_update.emit(row_index, str(e)[:100])
            return index, row_index, "错误", str(e)

    def run(self):
        """Main worker run method - uses ThreadPoolExecutor for parallel processing
        
        Key improvements:
        1. Each task updates UI immediately upon completion (not batched)
        2. Stop is processed immediately - we don't wait for running tasks
        3. Running tasks complete in background, their results are already sent to UI
        """
        import time
        from concurrent.futures import ThreadPoolExecutor
        
        was_stopped = False
        executor = None
        
        try:
            if not os.path.exists(self.output_dir):
                os.makedirs(self.output_dir)

            # Create executor (not using 'with' to allow non-blocking shutdown)
            executor = ThreadPoolExecutor(max_workers=self.max_workers)
            futures = {}
            
            # Submit tasks with delay between each
            for i, filepath in enumerate(self.files):
                # Check stop before submitting new task
                if self._stop_event.is_set():
                    # Mark all remaining as stopped
                    for j in range(i, len(self.files)):
                        row_idx = self.row_indices[j] if j < len(self.row_indices) else j
                        self.signals.status_update.emit(row_idx, "已终止")
                        self.signals.result_update.emit(row_idx, "用户已终止任务")
                    was_stopped = True
                    break
                
                future = executor.submit(self.process_single_file, i, filepath)
                futures[future] = i
                
                # Add delay between task submissions (except for last one)
                # Use short sleep intervals to check stop_event more frequently
                if i < len(self.files) - 1:
                    # First 2-second delay (reduced from 5s for better throughput)
                    for _ in range(20):  # 2 seconds total, check every 0.1s
                        if self._stop_event.is_set():
                            break
                        time.sleep(0.1)
                    
                    # Check if previous task is still processing
                    # If still processing, wait up to 3 more seconds
                    if not self._stop_event.is_set():
                        prev_future = list(futures.keys())[-1] if futures else None
                        if prev_future and not prev_future.done():
                            # Previous task still running, add up to 3 second delay
                            for _ in range(30):  # 3 seconds max
                                if self._stop_event.is_set():
                                    break
                                if prev_future.done():
                                    break  # Task completed, proceed immediately
                                time.sleep(0.1)
            
            # If stop was triggered during submission
            if self._stop_event.is_set():
                was_stopped = True
                # Mark any unstarted tasks as stopped
                for i in range(len(self.files)):
                    row_idx = self.row_indices[i] if i < len(self.row_indices) else i
                    # Only update if still showing pending/processing
                    # (completed tasks already updated themselves)
                self._finish_immediately(executor, futures)
                return
            
            # Wait for remaining futures with very short timeout
            # This allows us to check stop_event frequently
            pending = set(futures.keys())
            while pending and not self._stop_event.is_set():
                # Check each pending future with short timeout
                done = set()
                for future in list(pending):
                    if future.done():
                        done.add(future)
                        # Result already sent to UI by process_single_file
                        try:
                            future.result(timeout=0)  # Just to catch exceptions
                        except:
                            pass
                
                pending -= done
                
                if pending and not self._stop_event.is_set():
                    time.sleep(0.1)  # Short sleep to avoid busy loop
            
            # If stop was triggered while waiting
            if self._stop_event.is_set():
                was_stopped = True
                # Update any tasks that are still pending/processing
                for future in pending:
                    idx = futures[future]
                    row_idx = self.row_indices[idx] if idx < len(self.row_indices) else idx
                    self.signals.status_update.emit(row_idx, "已终止")
                    self.signals.result_update.emit(row_idx, "用户已终止任务")

        except Exception as e:
            print(f"Fatal batch error: {e}")
            import traceback
            traceback.print_exc()
        finally:
            # Shutdown executor without waiting (let running tasks finish in background)
            if executor:
                executor.shutdown(wait=False, cancel_futures=True)
            
            # Emit appropriate signal
            if was_stopped or self._stop_event.is_set():
                self.signals.stopped.emit()
            else:
                self.signals.finished.emit()
    
    def _finish_immediately(self, executor, futures):
        """Helper to finish immediately when stopped during submission"""
        if executor:
            executor.shutdown(wait=False, cancel_futures=True)

    def save_results(self, filepath, text, markdown_text, tables_html, layout_data, raw_result):
        base_name = os.path.splitext(os.path.basename(filepath))[0]
        fmt = self.export_fmt.lower()
        
        def p(ext): return os.path.join(self.output_dir, f"{base_name}.{ext}")

        try:
            # ========== PP-OCRv5 Formats ==========
            # Note: fmt is lowercased, so we check for lowercase versions
            if fmt == "txt":
                with open(p("txt"), "w", encoding="utf-8") as f:
                    f.write(text)
                return f"已导出: {base_name}.txt"
            
            elif fmt == "json":
                with open(p("json"), "w", encoding="utf-8") as f:
                    output = {"file": os.path.basename(filepath), "text": text}
                    if raw_result:
                        output["raw_result"] = raw_result
                    json.dump(output, f, ensure_ascii=False, indent=2)
                return f"已导出: {base_name}.json"
            
            elif "csv" in fmt:
                with open(p("csv"), "w", newline="", encoding="utf-8") as f:
                    writer = csv.writer(f)
                    writer.writerow(["行号", "文本"])
                    for idx, line in enumerate(text.split("\n")):
                        if line.strip():
                            writer.writerow([idx + 1, line])
                return f"已导出: {base_name}.csv"
            
            elif "可搜索pdf" in fmt:
                with open(p("txt"), "w", encoding="utf-8") as f:
                    f.write(text)
                return f"已导出: {base_name}.txt (可搜索PDF开发中)"
            
            elif "带标注" in fmt:
                with open(p("txt"), "w", encoding="utf-8") as f:
                    f.write(text)
                return f"已导出: {base_name}.txt (标注图片开发中)"

            # ========== PP-StructureV3 Formats ==========
            elif "docx" in fmt or "word" in fmt:
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading(base_name, 0)
                    
                    content = markdown_text if markdown_text else text
                    
                    for line in content.split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        if line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("### "):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith("- ") or line.startswith("* "):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        else:
                            doc.add_paragraph(line)
                    
                    doc.save(p("docx"))
                    return f"已导出: {base_name}.docx"
                except ImportError:
                    with open(p("md"), "w", encoding="utf-8") as f:
                        f.write(f"# {base_name}\n\n{markdown_text or text}")
                    return f"导出失败: 需要安装python-docx (已保存为MD)"
                except Exception as e:
                    with open(p("md"), "w", encoding="utf-8") as f:
                        f.write(f"# {base_name}\n\n{markdown_text or text}")
                    return f"DOCX导出失败: {str(e)[:30]}"
            
            elif "xlsx" in fmt or "excel" in fmt:
                try:
                    import openpyxl
                    wb = openpyxl.Workbook()
                    table_count = 0
                    
                    # First try to use HTML tables from PP-StructureV3
                    if tables_html:
                        for idx, html_table in enumerate(tables_html):
                            parsed_tables = parse_html_tables(html_table)
                            for t_idx, table_data in enumerate(parsed_tables):
                                if table_count == 0:
                                    ws = wb.active
                                    ws.title = f"表格1"
                                else:
                                    ws = wb.create_sheet(title=f"表格{table_count + 1}")
                                
                                for row_idx, row in enumerate(table_data, 1):
                                    for col_idx, cell in enumerate(row, 1):
                                        ws.cell(row=row_idx, column=col_idx, value=cell)
                                table_count += 1
                        
                        wb.save(p("xlsx"))
                        return f"已导出: {base_name}.xlsx ({table_count}个表格)"
                    
                    # Try to parse markdown tables (now returns list of tables)
                    elif markdown_text:
                        md_tables = parse_markdown_table(markdown_text)
                        if md_tables:
                            for t_idx, table_data in enumerate(md_tables):
                                if t_idx == 0:
                                    ws = wb.active
                                    ws.title = f"表格1"
                                else:
                                    ws = wb.create_sheet(title=f"表格{t_idx + 1}")
                                
                                for row_idx, row in enumerate(table_data, 1):
                                    for col_idx, cell in enumerate(row, 1):
                                        ws.cell(row=row_idx, column=col_idx, value=cell)
                            
                            wb.save(p("xlsx"))
                            return f"已导出: {base_name}.xlsx ({len(md_tables)}个表格)"
                    
                    # Fallback: use layout data or plain text as structured data
                    ws = wb.active
                    ws.title = "识别结果"
                    if layout_data:
                        ws.append(["序号", "类型", "内容"])
                        for idx, block in enumerate(layout_data):
                            ws.append([idx + 1, block.get("label", "text"), block.get("content", "")])
                    else:
                        ws.append(["行号", "文本"])
                        for idx, line in enumerate(text.split("\n")):
                            if line.strip():
                                ws.append([idx + 1, line])
                    
                    wb.save(p("xlsx"))
                    return f"已导出: {base_name}.xlsx"
                except ImportError:
                    with open(p("csv"), "w", newline="", encoding="utf-8") as f:
                        writer = csv.writer(f)
                        writer.writerow(["行号", "文本"])
                        for idx, line in enumerate(text.split("\n")):
                            writer.writerow([idx + 1, line])
                    return f"导出失败: 需要安装openpyxl (已保存为CSV)"
                except Exception as e:
                    return f"Excel导出失败: {str(e)[:30]}"
            
            elif "markdown" in fmt or "md" in fmt:
                content = markdown_text if markdown_text else text
                with open(p("md"), "w", encoding="utf-8") as f:
                    if not content.startswith("#"):
                        f.write(f"# {base_name}\n\n")
                    f.write(content)
                return f"已导出: {base_name}.md"
            
            elif "html" in fmt:
                # If we have HTML tables, combine them properly
                if tables_html:
                    html_content = f"""<!DOCTYPE html>
<html lang="zh">
<head>
    <meta charset="utf-8">
    <title>{base_name}</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
    </style>
</head>
<body>
    <h1>{base_name}</h1>
    {'<hr>'.join(tables_html)}
</body>
</html>"""
                    with open(p("html"), "w", encoding="utf-8") as f:
                        f.write(html_content)
                    return f"已导出: {base_name}.html"
                
                # Fallback: convert markdown to HTML
                content = markdown_text if markdown_text else text
                html_lines = []
                for line in content.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("# "):
                        html_lines.append(f"<h1>{line[2:]}</h1>")
                    elif line.startswith("## "):
                        html_lines.append(f"<h2>{line[3:]}</h2>")
                    elif line.startswith("### "):
                        html_lines.append(f"<h3>{line[4:]}</h3>")
                    else:
                        html_lines.append(f"<p>{line}</p>")
                
                html_content = f"""<!DOCTYPE html>
<html lang="zh">
<head>
    <meta charset="utf-8">
    <title>{base_name}</title>
    <style>body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; }}</style>
</head>
<body>{''.join(html_lines)}</body>
</html>"""
                with open(p("html"), "w", encoding="utf-8") as f:
                    f.write(html_content)
                return f"已导出: {base_name}.html"
            
            elif "版面树" in fmt or "layout" in fmt:
                # This format was removed in favor of just "JSON" for PP-StructureV3
                # But we keep backward compatibility
                with open(p("json"), "w", encoding="utf-8") as f:
                    output = {
                        "file": os.path.basename(filepath),
                        "layout_data": layout_data if layout_data else [],
                        "raw_result": raw_result
                    }
                    json.dump(output, f, ensure_ascii=False, indent=2)
                return f"已导出: {base_name}.json"
            
            elif "pdf" in fmt and "可搜索" not in fmt:
                try:
                    from reportlab.lib.pagesizes import A4
                    from reportlab.pdfgen import canvas
                    from reportlab.pdfbase import pdfmetrics
                    from reportlab.pdfbase.ttfonts import TTFont
                    from reportlab.lib.units import cm
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.enums import TA_LEFT
                    
                    content = markdown_text if markdown_text else text
                    pdf_path = p("pdf")
                    
                    # Try to register Chinese font
                    try:
                        pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                        font_name = 'SimSun'
                    except:
                        try:
                            pdfmetrics.registerFont(TTFont('MSYaHei', 'C:/Windows/Fonts/msyh.ttc'))
                            font_name = 'MSYaHei'
                        except:
                            font_name = 'Helvetica'
                    
                    # Create document with proper margins
                    doc = SimpleDocTemplate(
                        pdf_path,
                        pagesize=A4,
                        leftMargin=2*cm,
                        rightMargin=2*cm,
                        topMargin=2*cm,
                        bottomMargin=2*cm
                    )
                    
                    # Define styles
                    styles = getSampleStyleSheet()
                    normal_style = ParagraphStyle(
                        'Normal_CN',
                        parent=styles['Normal'],
                        fontName=font_name,
                        fontSize=12,
                        leading=18,
                        wordWrap='CJK'
                    )
                    h1_style = ParagraphStyle(
                        'H1_CN',
                        parent=styles['Heading1'],
                        fontName=font_name,
                        fontSize=18,
                        leading=24,
                        spaceAfter=12
                    )
                    h2_style = ParagraphStyle(
                        'H2_CN',
                        parent=styles['Heading2'],
                        fontName=font_name,
                        fontSize=16,
                        leading=20,
                        spaceAfter=10
                    )
                    h3_style = ParagraphStyle(
                        'H3_CN',
                        parent=styles['Heading3'],
                        fontName=font_name,
                        fontSize=14,
                        leading=18,
                        spaceAfter=8
                    )
                    
                    # Build story (content)
                    story = []
                    
                    for line in content.split("\n"):
                        line = line.strip()
                        if not line:
                            story.append(Spacer(1, 6))
                            continue
                        
                        # Escape XML special chars
                        line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        
                        # Handle headings
                        if line.startswith("# "):
                            story.append(Paragraph(line[2:], h1_style))
                        elif line.startswith("## "):
                            story.append(Paragraph(line[3:], h2_style))
                        elif line.startswith("### "):
                            story.append(Paragraph(line[4:], h3_style))
                        else:
                            story.append(Paragraph(line, normal_style))
                    
                    doc.build(story)
                    return f"已导出: {base_name}.pdf"
                except ImportError as ie:
                    content = markdown_text if markdown_text else text
                    with open(p("md"), "w", encoding="utf-8") as f:
                        if not content.startswith("#"):
                            f.write(f"# {base_name}\n\n")
                        f.write(content)
                    return f"导出失败: 需要安装reportlab (已保存为MD)"
                except Exception as e:
                    with open(p("md"), "w", encoding="utf-8") as f:
                        content = markdown_text if markdown_text else text
                        f.write(f"# {base_name}\n\n{content}")
                    return f"PDF导出失败: {str(e)[:30]}"

            # ========== PaddleOCR-VL Formats ==========
            elif "latex" in fmt or "公式" in fmt:
                with open(p("tex"), "w", encoding="utf-8") as f:
                    f.write(f"% {base_name}\n% LaTeX export\n\n{text}")
                return f"已导出: {base_name}.tex"
            
            elif "语义" in fmt or "kvp" in fmt:
                # This format was removed in favor of just "JSON" for PaddleOCR-VL
                with open(p("json"), "w", encoding="utf-8") as f:
                    json.dump({"file": os.path.basename(filepath), "semantic_kvp": text}, f, ensure_ascii=False, indent=2)
                return f"已导出: {base_name}.json"
            
            elif "叙述" in fmt or "narrative" in fmt:
                # This format was removed in favor of just "TXT" for PaddleOCR-VL
                with open(p("txt"), "w", encoding="utf-8") as f:
                    f.write(text)
                return f"已导出: {base_name}.txt"
            
            elif "代码" in fmt or "code" in fmt:
                with open(p("py"), "w", encoding="utf-8") as f:
                    f.write(f"# Extracted from {base_name}\n\n{text}")
                return f"已导出: {base_name}.py"
            
            else:
                with open(p("txt"), "w", encoding="utf-8") as f:
                    f.write(text)
                return f"已导出: {base_name}.txt"
                
        except Exception as e:
            return f"导出错误: {str(e)[:50]}"


class BatchProcessor:
    def __init__(self):
        self.thread_pool = QThreadPool.globalInstance()
        self.current_worker = None

    def process(self, mode, file_list, model_override=None, output_dir=None, export_fmt=None, max_workers=5, row_indices=None):
        if not output_dir:
            output_dir = config.get_batch_config(mode).get("output_dir", "output")
        if not export_fmt:
            export_fmt = "TXT"

        self.current_worker = BatchWorker(mode, file_list, model_override, output_dir, export_fmt, max_workers, row_indices)
        return self.current_worker

    def start(self, worker):
        self.thread_pool.start(worker)

    def stop(self):
        if self.current_worker:
            self.current_worker.stop()


batch_processor = BatchProcessor()
